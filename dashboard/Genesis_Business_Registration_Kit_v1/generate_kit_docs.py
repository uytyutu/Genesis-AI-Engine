#!/usr/bin/env python3
"""Generate PDF and DOCX from Genesis Business Registration Kit v1 markdown sources."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

KIT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = KIT_DIR / "output"

SOURCES = [
    ("00_README_Kit_Anleitung.md", "00_README"),
    ("01_GewA1_Felder_Antworten.md", "01_GewA1"),
    ("02_Taetigkeitsbeschreibung_Genesis.md", "02_Taetigkeit"),
    ("03_Unternehmensbeschreibung_Genesis.md", "03_Unternehmen"),
    ("04_Jobcenter_Selbststaendigkeit_Anziege.md", "04_Jobcenter"),
    ("05_Impressum_Vorlage.md", "05_Impressum"),
    ("06_Datenschutzerklaerung_Vorlage.md", "06_Datenschutz"),
    ("07_AGB_Vorlage.md", "07_AGB"),
    ("08_Checkliste_Gewerbe_und_Launch.md", "08_Checkliste"),
]


def read_source(name: str) -> str:
    return (KIT_DIR / name).read_text(encoding="utf-8")


def strip_md_for_plain(text: str) -> str:
    """Light markdown cleanup for PDF plain text."""
    lines = []
    in_code = False
    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code = not in_code
            if in_code:
                lines.append("---")
            else:
                lines.append("---")
            continue
        if in_code:
            lines.append(line)
            continue
        line = re.sub(r"^#+\s*", "", line)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"`([^`]+)`", r"\1", line)
        lines.append(line)
    return "\n".join(lines)


def latin1_safe(text: str) -> str:
    """FPDF core fonts are latin-1; replace unsupported chars."""
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u201e": '"',
        "\u201c": '"',
        "\u2018": "'",
        "\u2019": "'",
        "\u2026": "...",
        "\u20ac": "EUR",
        "\u2611": "[x]",
        "\u2610": "[ ]",
        "\u2192": "->",
        "\u2265": ">=",
        "\u00b7": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class KitPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 8, "Genesis Business Registration Kit v1", align="R")
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Seite {self.page_no()}", align="C")

    def write_body(self, text: str):
        self.set_font("Helvetica", "", 10)
        effective_w = self.w - self.l_margin - self.r_margin
        for line in text.splitlines():
            if not line.strip():
                self.ln(4)
                continue
            if re.match(r"^\|[-\s|]+\|$", line.strip()):
                continue
            safe = latin1_safe(line)
            if line.startswith("## "):
                self.ln(3)
                self.set_font("Helvetica", "B", 12)
                self.multi_cell(effective_w, 6, safe[3:].strip() if safe.startswith("##") else safe.lstrip("# ").strip())
                self.set_font("Helvetica", "", 10)
                continue
            if line.startswith("### "):
                self.ln(2)
                self.set_font("Helvetica", "B", 11)
                self.multi_cell(effective_w, 6, safe.lstrip("# ").strip())
                self.set_font("Helvetica", "", 10)
                continue
            if line.startswith("# "):
                self.set_font("Helvetica", "B", 14)
                self.multi_cell(effective_w, 7, safe.lstrip("# ").strip())
                self.set_font("Helvetica", "", 10)
                self.ln(2)
                continue
            if len(safe) > 200:
                self.set_font("Helvetica", "", 9)
            self.multi_cell(effective_w, 5, safe)
            self.set_font("Helvetica", "", 10)


def build_combined_pdf() -> Path:
    pdf = KitPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    for src, _ in SOURCES:
        text = strip_md_for_plain(read_source(src))
        pdf.write_body(text)
        pdf.add_page()
    out = OUTPUT_DIR / "Genesis_Business_Registration_Kit_v1.pdf"
    pdf.output(str(out))
    return out


def build_individual_pdfs() -> list[Path]:
    paths = []
    for src, short in SOURCES:
        pdf = KitPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.write_body(strip_md_for_plain(read_source(src)))
        out = OUTPUT_DIR / f"{short}.pdf"
        pdf.output(str(out))
        paths.append(out)
    return paths


def add_md_to_docx(doc: Document, text: str) -> None:
    in_code = False
    for line in text.splitlines():
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(line)
            p.style = "Intense Quote"
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip().startswith("|") and "---" not in line:
            doc.add_paragraph(line.strip())
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif not line.strip():
            doc.add_paragraph("")
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            doc.add_paragraph(clean)


def build_combined_docx() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_heading("Genesis Business Registration Kit v1", 0)
    doc.add_paragraph(
        "Etappe 1: Vorlagen ohne persönliche Daten. "
        "Ersetzen Sie alle {{PLATZHALTER}} vor der Einreichung."
    )
    for src, _ in SOURCES:
        doc.add_page_break()
        add_md_to_docx(doc, read_source(src))
    out = OUTPUT_DIR / "Genesis_Business_Registration_Kit_v1.docx"
    doc.save(str(out))
    return out


def build_individual_docx() -> list[Path]:
    paths = []
    for src, short in SOURCES:
        doc = Document()
        add_md_to_docx(doc, read_source(src))
        out = OUTPUT_DIR / f"{short}.docx"
        doc.save(str(out))
        paths.append(out)
    return paths


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf = build_combined_pdf()
    docx = build_combined_docx()
    ind_pdf = build_individual_pdfs()
    ind_docx = build_individual_docx()
    print(f"Generated: {pdf}")
    print(f"Generated: {docx}")
    print(f"Individual PDFs: {len(ind_pdf)}")
    print(f"Individual DOCX: {len(ind_docx)}")


if __name__ == "__main__":
    main()
